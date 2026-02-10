import os
import pandas as pd
import torch
from docx import Document
from tqdm import tqdm

# Qdrant & Embedding
from qdrant_client import QdrantClient, models
from sentence_transformers import SentenceTransformer

# LLM & Reranker
from openai import OpenAI
from transformers import AutoTokenizer, AutoModelForCausalLM

# DeepEval 相關
from deepeval import evaluate
from deepeval.metrics import (
    FaithfulnessMetric,
    AnswerRelevancyMetric,
    ContextualRecallMetric,
    ContextualPrecisionMetric,
    ContextualRelevancyMetric
)
from deepeval.test_case import LLMTestCase
from deepeval.models import DeepEvalBaseLLM

# ==========================================
# 1. 設定區 (Configuration)
# ==========================================

# Qdrant 設定
qdrant_client = QdrantClient(host="localhost", port=6333)
COLLECTION_NAME = "day6_hw_qwen"

# Embedding 模型
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
VECTOR_SIZE = 384

# LLM 設定 (一般生成用)
llm_client = OpenAI(
    base_url="https://ws-03.wade0426.me/v1", 
    api_key="sk-dummy-key"
)
LLM_MODEL_NAME = "/models/gpt-oss-120b"
os.environ["OPENAI_API_KEY"] = "sk-dummy-key" # DeepEval 預設檢查用

# ==========================================
# 2. DeepEval 自定義模型 Wrapper
# ==========================================
class CustomOpenAIModel(DeepEvalBaseLLM):
    """
    讓 DeepEval 使用我們自架的 API (ws-03)
    """
    def __init__(self, model_name, base_url, api_key):
        self.model_name = model_name
        self.base_url = base_url
        self.api_key = api_key

    def load_model(self):
        return OpenAI(base_url=self.base_url, api_key=self.api_key)

    def generate(self, prompt: str) -> str:
        client = self.load_model()
        try:
            res = client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
            )
            return res.choices[0].message.content
        except Exception as e:
            return f"Error: {e}"

    async def a_generate(self, prompt: str) -> str:
        return self.generate(prompt)

    def get_model_name(self):
        return "Custom GPT-OSS-120B"

# 初始化 DeepEval 專用的模型物件
deep_eval_model = CustomOpenAIModel(
    model_name=LLM_MODEL_NAME,
    base_url="https://ws-03.wade0426.me/v1",
    api_key="sk-dummy-key"
)

# ==========================================
# 3. 載入 Qwen3-Reranker 模型
# ==========================================
model_path = "/home/kiwi/test_folder/day6/Qwen3-Reranker-0.6B"

print(f"正在載入 Reranker 模型: {model_path} ...")
if not os.path.exists(model_path):
    print(f"錯誤：找不到模型路徑 {model_path}，請確認路徑是否正確。")
    exit()

try:
    reranker_tokenizer = AutoTokenizer.from_pretrained(
        model_path, local_files_only=True, trust_remote_code=True
    )
    reranker_model = AutoModelForCausalLM.from_pretrained(
        model_path, local_files_only=True, trust_remote_code=True,
        dtype=torch.bfloat16 if torch.cuda.is_available() else torch.float32
    ).eval()
    
    if torch.cuda.is_available():
        reranker_model = reranker_model.to("cuda")
        print("Reranker 已載入至 GPU")
    else:
        print("警告: 未偵測到 GPU，Reranker 將使用 CPU 執行")
except Exception as e:
    print(f"Reranker 模型載入失敗: {e}")
    exit()

token_false_id = reranker_tokenizer.convert_tokens_to_ids("no")
token_true_id = reranker_tokenizer.convert_tokens_to_ids("yes")
max_reranker_length = 2048

prefix = "<|im_start|>system\nJudge whether the Document meets the requirements based on the Query and the Instruct provided. Note that the answer can only be \"yes\" or \"no\".<|im_end|>\n<|im_start|>user\n"
suffix = "<|im_end|>\n<|im_start|>assistant\n"
prefix_tokens = reranker_tokenizer.encode(prefix, add_special_tokens=False)
suffix_tokens = reranker_tokenizer.encode(suffix, add_special_tokens=False)

# ==========================================
# 4. 核心函式 (Rerank, Search, Generation)
# ==========================================
def format_instruction(instruction, query, doc):
    if instruction is None: instruction = '根據查詢檢索相關文件'
    return f"<Instruct>: {instruction}\n<Query>: {query}\n<Document>: {doc}"

@torch.no_grad()
def get_rerank_scores(query, candidate_docs):
    if not candidate_docs: return []
    input_texts = [format_instruction(None, query, doc) for doc in candidate_docs]
    batch_input_ids = []
    for text in input_texts:
        text_ids = reranker_tokenizer.encode(text, add_special_tokens=False, truncation=True, max_length=max_reranker_length - len(prefix_tokens) - len(suffix_tokens))
        full_ids = prefix_tokens + text_ids + suffix_tokens
        batch_input_ids.append(full_ids)

    max_len = max(len(ids) for ids in batch_input_ids)
    padded_input_ids = [ids + [reranker_tokenizer.pad_token_id] * (max_len - len(ids)) for ids in batch_input_ids]
    attention_masks = [[1] * len(ids) + [0] * (max_len - len(ids)) for ids in batch_input_ids]

    input_tensor = torch.tensor(padded_input_ids).to(reranker_model.device)
    mask_tensor = torch.tensor(attention_masks).to(reranker_model.device)
    outputs = reranker_model(input_ids=input_tensor, attention_mask=mask_tensor)
    last_token_indices = mask_tensor.sum(1) - 1
    target_logits = outputs.logits[torch.arange(outputs.logits.shape[0]), last_token_indices, :]
    scores_pair = torch.stack([target_logits[:, token_false_id], target_logits[:, token_true_id]], dim=1)
    probs = torch.nn.functional.softmax(scores_pair, dim=1)
    return probs[:, 1].tolist()

def hybrid_search_rerank(query, top_n=10, top_k=3):
    query_vector = embedding_model.encode([query])[0].tolist()
    search_result = qdrant_client.query_points(
        collection_name=COLLECTION_NAME,
        prefetch=[
            models.Prefetch(query=models.Document(text=query, model="Qdrant/bm25"), using="sparse", limit=top_n),
            models.Prefetch(query=query_vector, using="dense", limit=top_n),
        ],
        query=models.FusionQuery(fusion=models.Fusion.RRF),
        limit=top_n,
    )
    candidate_docs = [hit.payload['text'] for hit in search_result.points]
    if not candidate_docs: return []
    scores = get_rerank_scores(query, candidate_docs)
    ranked_results = sorted(zip(candidate_docs, scores), key=lambda x: x[1], reverse=True)
    return [doc for doc, score in ranked_results[:top_k]]

def query_rewrite(query, chat_history):
    history_text = "\n".join([f"{msg['role']}: {msg['content']}" for msg in chat_history[-4:]])
    system_prompt = f"你是 RAG 專家。將使用者問題改寫為搜尋語句。只輸出改寫後的句子。\n[對話歷史]\n{history_text}"
    try:
        res = llm_client.chat.completions.create(model=LLM_MODEL_NAME, messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": query}], temperature=0.3)
        return res.choices[0].message.content.strip()
    except: return query

def generate_answer(query, context, chat_history):
    system_prompt = f"你是 AI 助理。請根據 [參考資料] 回答。資料不足請回答無法回答。\n[參考資料]\n{context}"
    messages = [{"role": "system", "content": system_prompt}] + chat_history[-4:] + [{"role": "user", "content": query}]
    try:
        res = llm_client.chat.completions.create(model=LLM_MODEL_NAME, messages=messages, temperature=0.7)
        return res.choices[0].message.content
    except Exception as e: return f"Error: {e}"

# ==========================================
# 5. 資料處理與評測流程
# ==========================================
def ingest_data():
    file_path = "qa_data.docx"
    if not os.path.exists(file_path): return
    doc = Document(file_path)
    texts = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 10]
    
    qdrant_client.recreate_collection(
        collection_name=COLLECTION_NAME,
        vectors_config={"dense": models.VectorParams(size=VECTOR_SIZE, distance=models.Distance.COSINE)},
        sparse_vectors_config={"sparse": models.SparseVectorParams(modifier=models.Modifier.IDF)}
    )
    dense_vectors = embedding_model.encode(texts).tolist()
    points = [models.PointStruct(id=i, vector={"dense": v, "sparse": models.Document(text=t, model="Qdrant/bm25")}, payload={"text": t}) for i, (t, v) in enumerate(zip(texts, dense_vectors))]
    qdrant_client.upsert(collection_name=COLLECTION_NAME, points=points)
    print(f"資料庫建立完成，共 {len(points)} 筆。")

def run_evaluation(question, actual_output, retrieval_context, expected_output):
    test_case = LLMTestCase(input=question, actual_output=actual_output, retrieval_context=retrieval_context, expected_output=expected_output)
    metrics = [
        FaithfulnessMetric(threshold=0.5, model=deep_eval_model, include_reason=False),
        AnswerRelevancyMetric(threshold=0.5, model=deep_eval_model, include_reason=False),
        ContextualRecallMetric(threshold=0.5, model=deep_eval_model, include_reason=False),
        ContextualPrecisionMetric(threshold=0.5, model=deep_eval_model, include_reason=False),
        ContextualRelevancyMetric(threshold=0.5, model=deep_eval_model, include_reason=False)
    ]
    names = ['Faithfulness', 'Answer_Relevancy', 'Contextual_Recall', 'Contextual_Precision', 'Contextual_Relevancy']
    results = {}
    for name, m in zip(names, metrics):
        try:
            m.measure(test_case)
            results[name] = m.score
        except: results[name] = 0.0
    return results

# ==========================================
# 6. 主程式 (只讀取前五筆)
# ==========================================
def main():
    ingest_data()
    
    input_xlsx = "day6_HW_questions.csv.xlsx"
    truth_xlsx = "questions_answer.csv.xlsx"
    output_csv = "day6_HW_questions_result.csv"
    
    if not (os.path.exists(input_xlsx) and os.path.exists(truth_xlsx)):
        print("錯誤：找不到 Excel 檔案")
        return

    # 讀取並只取前 5 筆
    df_questions = pd.read_excel(input_xlsx).head(5).copy()
    df_truth = pd.read_excel(truth_xlsx)
    truth_map = dict(zip(df_truth['q_id'], df_truth['answer']))

    # 預設欄位與型別優化 (避免 FutureWarning)
    for col in ['answer', 'Faithfulness', 'Answer_Relevancy', 'Contextual_Recall', 'Contextual_Precision', 'Contextual_Relevancy']:
        df_questions[col] = ""
        df_questions[col] = df_questions[col].astype(object)

    print(f"開始處理前 {len(df_questions)} 題...")
    chat_history = [] 

    for index, row in tqdm(df_questions.iterrows(), total=len(df_questions)):
        q_id, question = row['q_id'], str(row['questions'])
        expected_output = str(truth_map.get(q_id, ""))
        
        rw_q = query_rewrite(question, chat_history)
        retrieved_docs = hybrid_search_rerank(rw_q)
        actual_output = generate_answer(question, "\n\n".join(retrieved_docs), chat_history)
        
        chat_history.append({"role": "user", "content": question})
        chat_history.append({"role": "assistant", "content": actual_output})
        if len(chat_history) > 4: chat_history = chat_history[-4:]

        print(f"\n[Q{q_id}] 評測中...")
        metrics = run_evaluation(question, actual_output, retrieved_docs, expected_output)
        
        df_questions.at[index, 'answer'] = actual_output
        for k, v in metrics.items():
            df_questions.at[index, k] = v

    df_questions.to_csv(output_csv, index=False, encoding='utf-8-sig')
    print(f"\n完成！結果儲存至 {output_csv}")

if __name__ == "__main__":
    main()